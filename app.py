import streamlit as st
import pandas as pd
import random
from docx import Document
from io import BytesIO

# -----------------------------
# CONFIG
# -----------------------------
st.set_page_config(page_title="UrbanLab System", layout="wide")

st.markdown("""
<div style="background-color:#1e3a8a;color:white;padding:12px;border-radius:8px;text-align:center;font-weight:bold;margin-bottom:20px;">
🏙️ URBAN POLICY SIMULATION SYSTEM
</div>
""", unsafe_allow_html=True)

st.title("UrbanLab — Social Disorganization Simulator")

st.markdown("""
1️⃣ Analyse real neighbourhood  
2️⃣ Build criminological diagnosis  
3️⃣ Design intervention  
4️⃣ Evaluate and exchange reports  
""")

st.divider()

# -----------------------------
# BARRIOS REALES
# -----------------------------
barrios_reales = [
    "Polígono Sur (Sevilla)",
    "El Raval (Barcelona)",
    "El Cabanyal (Valencia)",
    "Puente de Vallecas (Madrid)",
    "Usera (Madrid)",
    "Ciutat Meridiana (Barcelona)",
    "La Mina (Sant Adrià del Besòs)"
]

if "barrio" not in st.session_state:
    st.session_state.barrio = random.choice(barrios_reales)

st.header("📍 Assigned neighbourhood")
st.success(st.session_state.barrio)

st.markdown("""
🔎 Use real sources (statistics, press, academic work) to build your diagnosis.
""")

# -----------------------------
# DIAGNÓSTICO
# -----------------------------
st.header("Step 1 — Criminological diagnosis")

diagnostico = st.text_area("""
Describe the neighbourhood using social disorganization theory:

- Social cohesion  
- Informal control  
- Structural conditions  
- Criminogenic factors  
""", height=200)

# -----------------------------
# PLAN
# -----------------------------
st.header("Step 2 — Intervention plan")

plan = st.text_area("Design your intervention", height=200)

# -----------------------------
# INTERPRETADOR
# -----------------------------
def interpretar_plan(plan):

    texto = plan.lower()

    categorias = {
        "Control formal": {
            "kw":["polic","vigilancia","cámaras"],
            "impacto":{"policia":15,"control":10},
            "tipo":"punitiva"
        },
        "Cohesión social": {
            "kw":["vecin","comunit","asociaciones"],
            "impacto":{"cohesion":15,"control":10},
            "tipo":"estructural"
        },
        "Urbanismo": {
            "kw":["urban","espacio","rehabilitación"],
            "impacto":{"desorganizacion":-15},
            "tipo":"estructural"
        },
        "Intervención económica": {
            "kw":["empleo","educa","formación"],
            "impacto":{"pobreza":-15},
            "tipo":"estructural"
        }
    }

    cambios = {"policia":0,"cohesion":0,"control":0,"desorganizacion":0,"pobreza":0}
    contribuciones = []
    tipos = set()

    for nombre,data in categorias.items():
        for kw in data["kw"]:
            if kw in texto:
                tipos.add(data["tipo"])
                for k,v in data["impacto"].items():
                    cambios[k]+=v
                contribuciones.append(nombre)
                break

    tipo_final = "mixta" if len(tipos)>1 else list(tipos)[0] if tipos else "indefinida"
    score = len(contribuciones)*2

    return cambios, contribuciones, tipo_final, min(score,10)

# -----------------------------
# EJECUCIÓN
# -----------------------------
if st.button("Run simulation"):

    if len(diagnostico) < 50:
        st.warning("Diagnosis too short")
        st.stop()

    if len(plan) < 50:
        st.warning("Plan too short")
        st.stop()

    cambios, contribuciones, tipo, score = interpretar_plan(plan)

    base = {
        "desorganizacion":70,
        "cohesion":40,
        "control":30,
        "pobreza":60,
        "policia":40
    }

    for k in cambios:
        if k in base:
            base[k]+=cambios[k]

    delito = (
        base["desorganizacion"]*0.4 +
        base["pobreza"]*0.3 -
        base["cohesion"]*0.3 -
        base["control"]*0.2
    )

    st.session_state.resultado = delito
    st.session_state.contrib = contribuciones
    st.session_state.tipo = tipo
    st.session_state.score = score

# -----------------------------
# RESULTADOS
# -----------------------------
if "resultado" in st.session_state:

    st.header("Step 3 — Results")

    c1,c2,c3 = st.columns(3)
    c1.metric("Crime level", round(st.session_state.resultado,2))
    c2.metric("Strategy", st.session_state.tipo)
    c3.metric("Score", st.session_state.score)

    df = pd.DataFrame(st.session_state.contrib, columns=["Interventions"])
    st.dataframe(df)

    # Feedback diagnóstico
    st.subheader("🧠 Diagnostic feedback")

    if "cohesion" not in diagnostico.lower():
        st.warning("Missing cohesion analysis")

    if "control" not in diagnostico.lower():
        st.warning("Missing informal control")

    if "pobre" not in diagnostico.lower():
        st.warning("Missing structural factors")

# -----------------------------
# INFORME WORD
# -----------------------------
if "resultado" in st.session_state:

    st.header("Step 4 — Generate report")

    if st.button("Generate report"):

        doc = Document()

        doc.add_heading('URBAN POLICY REPORT', 1)

        doc.add_heading('Neighbourhood',2)
        doc.add_paragraph(st.session_state.barrio)

        doc.add_heading('Diagnosis',2)
        doc.add_paragraph(diagnostico)

        doc.add_heading('Intervention',2)
        doc.add_paragraph(plan)

        doc.add_heading('Results',2)
        doc.add_paragraph(f"Crime level: {round(st.session_state.resultado,2)}")
        doc.add_paragraph(f"Strategy: {st.session_state.tipo}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "📄 Download report",
            buffer,
            file_name="urban_report.docx"
        )

# -----------------------------
# DRIVE + EVALUACIÓN
# -----------------------------
st.divider()
st.header("📄 Report exchange & evaluation")

grupo = st.selectbox("Group", ["Group A","Group B","Group C"])

links = {
    "Group A":{"upload":"LINK_A","review":"LINK_B"},
    "Group B":{"upload":"LINK_B","review":"LINK_C"},
    "Group C":{"upload":"LINK_C","review":"LINK_A"}
}

st.subheader("Upload your report")
st.markdown(f"[Open folder]({links[grupo]['upload']})")

st.subheader("Review reports")
st.markdown(f"[Open reports]({links[grupo]['review']})")

# -----------------------------
# RÚBRICA
# -----------------------------
st.header("Peer evaluation")

c1,c2 = st.columns(2)

with c1:
    coherencia = st.slider("Theoretical coherence",0,10,5)
    analisis = st.slider("Quality of diagnosis",0,10,5)

with c2:
    viabilidad = st.slider("Feasibility",0,10,5)
    innovacion = st.slider("Innovation",0,10,5)

nota = (coherencia+analisis+viabilidad+innovacion)/4
st.metric("Final grade", round(nota,2))

comentario = st.text_area("Comment")

# -----------------------------
# EXPORTAR EVALUACIÓN
# -----------------------------
if st.button("Generate evaluation report"):

    doc = Document()

    doc.add_heading('PEER EVALUATION REPORT',1)

    doc.add_paragraph(f"Group: {grupo}")
    doc.add_paragraph(f"Final grade: {round(nota,2)}")

    doc.add_heading("Scores",2)
    doc.add_paragraph(f"Coherence: {coherencia}")
    doc.add_paragraph(f"Diagnosis: {analisis}")
    doc.add_paragraph(f"Feasibility: {viabilidad}")
    doc.add_paragraph(f"Innovation: {innovacion}")

    doc.add_heading("Comment",2)
    doc.add_paragraph(comentario)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        "📥 Download evaluation",
        buffer,
        file_name="evaluation.docx"
    )
